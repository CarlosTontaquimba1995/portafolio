from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
PUBLIC = ROOT / "public"
DOWNLOADS = Path.home() / "Downloads"
PORTFOLIO_URL = "https://portafolio-66g7.onrender.com/"
LINKEDIN_URL = (
    "https://www.linkedin.com/in/carlos-vicente-tontaquimba-quinchuqui-3438b3384"
)
GITHUB_URL = "https://github.com/CarlosTontaquimba1995"
EMAIL = "carlos.tontaquimba1995@gmail.com"
PHONE = "+593 939618855"

EN = {
    "filename": "Carlos-Vicente-Tontaquimba-CV-EN.docx",
    "public_name": "cv-en.docx",
    "portfolio_label": "Portfolio",
    "profile_title": "Professional Profile",
    "profile": (
        "Computer Systems Engineer and Fullstack Developer with more than 6 years of "
        "specialized experience designing, developing, and implementing scalable "
        "microservice and microfrontend architectures. Strong command of the Java "
        "ecosystem (Spring Boot, Spring Cloud) and PHP (Laravel), combined with "
        "modern frontend development in Angular. Proven experience leading critical "
        "projects in the public and private sectors, implementing distributed "
        "solutions, asynchronous messaging (Kafka), serverless environments (AWS "
        "Lambda), and performance optimization (JMeter). Professional focused on "
        "agile methodologies, continuous delivery, and software engineering best "
        "practices. Available for international relocation under visa sponsorship "
        "in Spain."
    ),
    "skills_title": "Technical Skills",
    "skills": [
        ("Programming languages", "Java, PHP, JavaScript, TypeScript."),
        ("Backend frameworks", "Spring Boot, Spring Cloud, Laravel (v11), NestJS."),
        ("Frontend frameworks", "Angular (v17), JSF, Laravel Blade."),
        ("Databases", "PostgreSQL, MySQL."),
        (
            "Infrastructure and tools",
            "AWS Lambda (Serverless), Apache Kafka, Docker, Git, GitHub, RedHat Linux servers, JBoss EAP, JMeter (load testing).",
        ),
        ("Methodologies", "Scrum, GitFlow, Continuous Integration (CI/CD)."),
    ],
    "experience_title": "Professional Experience",
    "jobs": [
        {
            "title": "Information Systems Analyst 2 | Council of the Judiciary (Ecuador) | October 2025 – Present",
            "bullets": [
                "Design and execution of methodological phases for the software lifecycle under robust Microservice architectures with Spring Boot and Microfrontends with Angular.",
                "Strategic participation in operational plans for the successful implementation and rollout of new critical modules in the judicial system.",
            ],
        },
        {
            "title": "Fullstack Developer Consultant (Project contract) | Ministry of Telecommunications (Ecuador) | July 2025 – September 2025",
            "bullets": [
                "End-to-end development and delivery of scalable microservices using Laravel, improving the efficiency of several ministerial applications.",
                "Design and optimization of relational database schemas with PostgreSQL.",
                "Construction of professional user interfaces using Laravel/Blade.",
            ],
        },
        {
            "title": "Software Development Projects Expert | Council of the Judiciary (Ecuador) | February 2024 – June 2025",
            "bullets": [
                "Responsible for evolutionary maintenance and development of corporate microservices based on Java and Spring Boot.",
                "Design and development of modern, dynamic, and responsive user interfaces using Angular 17.",
                "Deployment and orchestration of microservices with Spring Cloud, ensuring a distributed and secure architecture on RedHat Linux servers.",
                "Implementation of automated load, stress, and performance tests with JMeter, ensuring high system availability.",
            ],
        },
        {
            "title": "Fullstack Software Engineer (Project contract) | Ministry of Economic and Social Inclusion (Ecuador) | August 2023 – December 2023",
            "bullets": [
                "Corrective and evolutionary maintenance of the critical agreements system for social protection and early childhood development programs.",
                'End-to-end development of the "COMITÉS" system, implementing backend microservices and modular Angular components for the frontend.',
            ],
        },
        {
            "title": "Software Development Analyst (Project contract) | Secretariat of Intercultural Bilingual Education (Ecuador) | February 2023 – July 2023",
            "bullets": [
                "Design, coding, and standardized implementation of technical requirements for the strategic KIPUJEA and CESLI systems.",
                "Execution of functional tests (unit and integration) to ensure code quality and stability before production.",
                "Preparation of detailed technical documentation of the developed architecture.",
            ],
        },
        {
            "title": "Associate Backend Developer | TRADE EC (Ecuador) | February 2022 – February 2023",
            "bullets": [
                "Development and optimization of efficient microservices using PHP (Laravel) and Node.js (NestJS).",
                "Creation of serverless services using AWS Lambda, reducing infrastructure costs and deployment times.",
                "Technical leadership of the critical API Gateway migration from Netflix Zuul (Java) to Amazon API Gateway.",
            ],
        },
        {
            "title": "Fullstack Application Developer | Latitude Foundation (Ecuador) | January 2019 – January 2022",
            "bullets": [
                "End-to-end development of electronic invoicing and grade-control systems using Spring Boot, Laravel, and Angular.",
                "Quality assurance through performance audits and stress tests with JMeter.",
            ],
        },
    ],
    "education_title": "Education",
    "education": [
        "Bachelor’s Degree in Computer Systems Engineering | Universidad Técnica del Norte (Ecuador), graduated in 2020.",
        "Associate Degree in Commerce and Administration (Informatics) | Instituto Tecnológico Otavalo (Ecuador), graduated in 2013.",
    ],
    "languages_title": "Languages",
    "languages": [
        "Spanish: Native.",
        "Kichwa: Native.",
        "English: B1 (Intermediate / professional technical).",
    ],
    "training_title": "Additional Training",
    "training": [
        "Professional Git and GitHub course.",
        "Design Thinking applied to product development.",
        "IT customer service and user support course.",
    ],
}

ES = {
    "filename": "Carlos-Vicente-Tontaquimba-CV-ES.docx",
    "public_name": "cv-es.docx",
    "portfolio_label": "Portafolio",
    "profile_title": "Perfil profesional",
    "profile": (
        "Ingeniero en Sistemas Computacionales y Desarrollador Fullstack con más de "
        "6 años de experiencia especializada en el diseño, desarrollo e "
        "implementación de arquitecturas de microservicios y microfrontends "
        "escalables. Sólido dominio del ecosistema Java (Spring Boot, Spring Cloud) "
        "y PHP (Laravel), combinado con desarrollo frontend moderno en Angular. "
        "Experiencia demostrada liderando proyectos críticos en el sector público y "
        "privado, implementando soluciones distribuidas, mensajería asíncrona "
        "(Kafka), entornos serverless (AWS Lambda) y optimización de rendimiento "
        "(JMeter). Profesional orientado a metodologías ágiles, entrega continua y "
        "buenas prácticas de ingeniería de software. Disponibilidad para "
        "reubicación internacional bajo patrocinio de visa en España."
    ),
    "skills_title": "Habilidades técnicas",
    "skills": [
        ("Lenguajes de programación", "Java, PHP, JavaScript, TypeScript."),
        ("Frameworks backend", "Spring Boot, Spring Cloud, Laravel (v11), NestJS."),
        ("Frameworks frontend", "Angular (v17), JSF, Laravel Blade."),
        ("Bases de datos", "PostgreSQL, MySQL."),
        (
            "Infraestructura y herramientas",
            "AWS Lambda (Serverless), Apache Kafka, Docker, Git, GitHub, servidores RedHat Linux, JBoss EAP, JMeter (pruebas de carga).",
        ),
        ("Metodologías", "Scrum, GitFlow, Integración Continua (CI/CD)."),
    ],
    "experience_title": "Experiencia profesional",
    "jobs": [
        {
            "title": "Analista de Sistemas de Información 2 | Consejo de la Judicatura (Ecuador) | Octubre 2025 – Actualidad",
            "bullets": [
                "Diseño y ejecución de fases metodológicas para el ciclo de vida del software bajo arquitecturas robustas de Microservicios con Spring Boot y Microfrontends con Angular.",
                "Participación estratégica en planes operativos para la implementación y despliegue exitoso de nuevos módulos críticos del sistema judicial.",
            ],
        },
        {
            "title": "Consultor Desarrollador Fullstack (Contrato por proyecto) | Ministerio de Telecomunicaciones (Ecuador) | Julio 2025 – Septiembre 2025",
            "bullets": [
                "Desarrollo y entrega llave en mano de microservicios escalables utilizando Laravel, mejorando la eficiencia de diversas aplicaciones ministeriales.",
                "Diseño y optimización de esquemas de bases de datos relacionales con PostgreSQL.",
                "Construcción de interfaces de usuario profesionales utilizando Laravel/Blade.",
            ],
        },
        {
            "title": "Experto en Proyectos de Desarrollo de Software | Consejo de la Judicatura (Ecuador) | Febrero 2024 – Junio 2025",
            "bullets": [
                "Responsable del mantenimiento evolutivo y desarrollo de microservicios corporativos basados en Java y Spring Boot.",
                "Diseño y desarrollo de interfaces de usuario modernas, dinámicas y responsivas utilizando Angular 17.",
                "Despliegue y orquestación de microservicios con Spring Cloud, asegurando una arquitectura distribuida y segura en servidores RedHat Linux.",
                "Implementación de pruebas automatizadas de carga, estrés y rendimiento con JMeter, garantizando alta disponibilidad del sistema.",
            ],
        },
        {
            "title": "Ingeniero de Software Fullstack (Contrato por proyecto) | Ministerio de Inclusión Económica y Social (Ecuador) | Agosto 2023 – Diciembre 2023",
            "bullets": [
                "Mantenimiento correctivo y evolutivo del sistema crítico de convenios para programas de protección social y desarrollo infantil.",
                'Desarrollo de extremo a extremo del sistema "COMITÉS", implementando microservicios backend y componentes modulares de Angular en el frontend.',
            ],
        },
        {
            "title": "Analista de Desarrollo de Software (Contrato por proyecto) | Secretaría de Educación Intercultural Bilingüe (Ecuador) | Febrero 2023 – Julio 2023",
            "bullets": [
                "Diseño, codificación e implementación estandarizada de requerimientos técnicos para los sistemas estratégicos KIPUJEA y CESLI.",
                "Ejecución de pruebas funcionales (unitarias e integración) para asegurar calidad y estabilidad del código antes de producción.",
                "Elaboración de documentación técnica detallada de la arquitectura desarrollada.",
            ],
        },
        {
            "title": "Desarrollador Backend Asociado | TRADE EC (Ecuador) | Febrero 2022 – Febrero 2023",
            "bullets": [
                "Desarrollo y optimización de microservicios eficientes utilizando PHP (Laravel) y Node.js (NestJS).",
                "Creación de servicios serverless con AWS Lambda, reduciendo costos de infraestructura y tiempos de despliegue.",
                "Liderazgo técnico de la migración crítica del API Gateway de Netflix Zuul (Java) a Amazon API Gateway.",
            ],
        },
        {
            "title": "Desarrollador de Aplicaciones Fullstack | Fundación Latitude (Ecuador) | Enero 2019 – Enero 2022",
            "bullets": [
                "Desarrollo de extremo a extremo de sistemas de facturación electrónica y control de calificaciones utilizando Spring Boot, Laravel y Angular.",
                "Aseguramiento de calidad mediante auditorías de rendimiento y pruebas de estrés con JMeter.",
            ],
        },
    ],
    "education_title": "Educación",
    "education": [
        "Ingeniería en Sistemas Computacionales | Universidad Técnica del Norte (Ecuador), graduado en 2020.",
        "Tecnología en Comercio y Administración (Informática) | Instituto Tecnológico Otavalo (Ecuador), graduado en 2013.",
    ],
    "languages_title": "Idiomas",
    "languages": [
        "Español: Nativo.",
        "Kichwa: Nativo.",
        "Inglés: B1 (Intermedio / técnico profesional).",
    ],
    "training_title": "Formación adicional",
    "training": [
        "Curso profesional de Git y GitHub.",
        "Design Thinking aplicado al desarrollo de producto.",
        "Curso de atención al cliente TI y soporte a usuarios.",
    ],
}


def set_run_font(run, *, size: int, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), "Calibri")
    if color is not None:
        run.font.color.rgb = color


def set_spacing(paragraph, *, before: int = 0, after: int = 4, line: float = 1.08) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "21")
    size_cs = OxmlElement("w:szCs")
    size_cs.set(qn("w:val"), "21")
    r_pr.extend([r_fonts, color, underline, size, size_cs])
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(r_pr)
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_labeled_link(document: Document, label: str, url: str) -> None:
    paragraph = document.add_paragraph()
    set_spacing(paragraph, after=1)
    run = paragraph.add_run(f"{label}: ")
    set_run_font(run, size=10.5)
    add_hyperlink(paragraph, url, url)


def add_section_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    set_spacing(paragraph, before=10, after=4)
    run = paragraph.add_run(text.upper())
    set_run_font(run, size=12, bold=True, color=RGBColor(0x11, 0x18, 0x27))
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "111827")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_body(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    set_spacing(paragraph, after=6)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5)


def add_bullet(document: Document, text: str, *, bold_prefix: str | None = None) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    set_spacing(paragraph, after=2)
    if bold_prefix:
        lead = paragraph.add_run(f"{bold_prefix}: ")
        set_run_font(lead, size=10.5, bold=True)
        rest = paragraph.add_run(text)
        set_run_font(rest, size=10.5)
        return
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5)


def build_cv(copy: dict) -> Document:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)

    styles = document.styles["Normal"]
    styles.font.name = "Calibri"
    styles.font.size = Pt(10.5)

    name = document.add_paragraph()
    set_spacing(name, after=4, line=1.0)
    name_run = name.add_run("CARLOS VICENTE TONTAQUIMBA QUINCHUQUI")
    set_run_font(name_run, size=18, bold=True)

    meta = document.add_paragraph()
    set_spacing(meta, after=1)
    meta_run = meta.add_run(f"Quito, Ecuador | {PHONE} | {EMAIL}")
    set_run_font(meta_run, size=10.5)

    add_labeled_link(document, copy["portfolio_label"], PORTFOLIO_URL)
    add_labeled_link(document, "LinkedIn", LINKEDIN_URL)
    add_labeled_link(document, "GitHub", GITHUB_URL)

    add_section_heading(document, copy["profile_title"])
    add_body(document, copy["profile"])

    add_section_heading(document, copy["skills_title"])
    for label, value in copy["skills"]:
        add_bullet(document, value, bold_prefix=label)

    add_section_heading(document, copy["experience_title"])
    for job in copy["jobs"]:
        title = document.add_paragraph()
        set_spacing(title, before=6, after=2)
        title_run = title.add_run(job["title"])
        set_run_font(title_run, size=10.5, bold=True)
        for bullet in job["bullets"]:
            add_bullet(document, bullet)

    add_section_heading(document, copy["education_title"])
    for item in copy["education"]:
        add_bullet(document, item)

    add_section_heading(document, copy["languages_title"])
    for item in copy["languages"]:
        add_bullet(document, item)

    add_section_heading(document, copy["training_title"])
    for item in copy["training"]:
        add_bullet(document, item)

    return document


def save_cv(copy: dict) -> None:
    document = build_cv(copy)
    PUBLIC.mkdir(exist_ok=True)
    public_path = PUBLIC / copy["public_name"]
    download_path = DOWNLOADS / copy["filename"]
    document.save(public_path)
    document.save(download_path)
    print(f"Wrote {public_path}")
    print(f"Wrote {download_path}")


if __name__ == "__main__":
    save_cv(EN)
    save_cv(ES)
